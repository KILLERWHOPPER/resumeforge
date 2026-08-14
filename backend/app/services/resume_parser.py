"""简历文件文本提取 — 支持 PDF、DOCX、TXT/Markdown"""

# ruff: noqa: TRY003

from __future__ import annotations

import io

from docx import Document
from pypdf import PdfReader

from app.core.exceptions import BadRequest

MIN_TEXT_CHARS = 50
MAX_TEXT_CHARS = 20000

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md", ".text", ".markdown")


def extract_resume_text(filename: str, content: bytes) -> str:
    """根据文件类型提取简历文本，返回清理后的纯文本"""
    name = filename.lower()
    if name.endswith(".pdf"):
        text = _extract_pdf(content)
    elif name.endswith(".docx"):
        text = _extract_docx(content)
    elif name.endswith((".txt", ".md", ".text", ".markdown")):
        text = _extract_plain_text(content)
    else:
        raise BadRequest("仅支持 PDF、DOCX、TXT/Markdown 文件")

    text = text.strip()
    if len(text) < MIN_TEXT_CHARS:
        raise BadRequest("未能从文件中识别出足够的文字内容，请检查文件是否为有效简历")
    return text[:MAX_TEXT_CHARS]


def _extract_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # noqa: BLE001 — 兼容各类损坏文件
        raise BadRequest("PDF 文件解析失败，请确认文件未损坏") from exc


def _extract_docx(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
        parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001 — 兼容各类损坏文件
        raise BadRequest("DOCX 文件解析失败，请确认文件未损坏") from exc


def _extract_plain_text(content: bytes) -> str:
    for encoding in ("utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")
